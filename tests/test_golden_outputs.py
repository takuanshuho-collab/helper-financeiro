"""Golden-master dos outputs .docx/.xlsx (T-2201, ADR-0019 M22).

Esta é a RÉGUA das refatorações T-2202 (`gerar_relatorio`) e T-2203
(`_aba_evolucao` e vizinhas): fixa o estado ATUAL dos geradores ANTES de
qualquer movimentação de código. Nenhum arquivo de produção muda nesta task —
o teste só gera os artefatos das fixtures do harness, extrai uma representação
canônica determinística e compara contra um JSON versionado em `tests/golden/`.

Extratores
----------
- ``extrair_docx``: percorre o corpo do documento na ORDEM real do XML
  (parágrafos e tabelas intercalados como aparecem) e emite ``[estilo, texto]``
  para cada parágrafo e para cada parágrafo de célula de tabela. Fixa prosa,
  ordem das seções e estilos — exatamente o que a suíte atual NÃO cobria e que
  a refatoração poderia alterar sem querer.
- ``extrair_xlsx``: por aba (na ordem do workbook), emite ``[coordenada, valor]``
  de cada célula não vazia. Célula com fórmula é lida com ``data_only=False`` e
  comparada como a STRING da fórmula (``"=SUM(...)"``), nunca como valor
  calculado — não há Excel no CI e o interesse é a fórmula em si.

O que o golden fixa deliberadamente: texto, estilo de parágrafo, ordem das
seções/abas, coordenadas e fórmulas. O que ele NÃO fixa (fora do contrato da
ADR): formatação de run (negrito/cor/fonte), objetos de gráfico, larguras de
coluna, merges e metadados do zip — justamente os pontos não-determinísticos
ou irrelevantes para a régua.

Máscaras (campo volátil → placeholder, SEMPRE no extrator, nunca no golden)
--------------------------------------------------------------------------
- ``<DATA>``: a capa do relatório imprime ``"Emitido em DD/MM/AAAA"`` a partir
  de ``date.today()`` (``outputs/relatorio.py``). Essa data muda a cada
  execução; sem máscara o golden quebraria todo dia. O extrator substitui
  QUALQUER ``dd/mm/aaaa`` por ``<DATA>``. É seguro: nenhuma outra string dos
  outputs usa esse formato (competências são ISO ``AAAA-MM`` e moeda é
  ``"R$ ..."``). A mesma máscara protege a data das cartas de proposta, caso
  venham a entrar na régua no futuro.

Regeneração
-----------
Regenera SÓ quando ``HF_REGENERAR_GOLDEN=1`` E fora de CI. Com ``CI`` setado a
regeneração é RECUSADA com falha explícita — a régua nunca se reescreve sozinha
no pipeline. Sem a flag, qualquer divergência falha com diff unificado legível.
"""
from __future__ import annotations

import difflib
import json
import os
import re
from collections.abc import Callable
from copy import deepcopy
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook

from contracts import PassoRoteiroIA, SecaoIA
from core.rubricas import Rubrica, serie_evolucao
from outputs.planilha import gerar_planilha
from outputs.relatorio import gerar_relatorio

GOLDEN_DIR = Path(__file__).parent / "golden"

# Data de geração no formato pt-BR (dd/mm/aaaa) — único campo volátil dos
# outputs. Mascarada no extrator para o golden não depender do dia da execução.
_RE_DATA_BR = re.compile(r"\d{2}/\d{2}/\d{4}")


def _mascarar(texto: str) -> str:
    """Neutraliza o campo volátil (data de geração) antes de congelar/comparar."""
    return _RE_DATA_BR.sub("<DATA>", texto)


# --------------------------------------------------------------- extratores
def _iter_blocos(doc: Any):
    """Gera parágrafos e tabelas do corpo na ordem exata do documento.

    ``doc.paragraphs`` e ``doc.tables`` vêm em listas separadas e perderiam a
    ordem em que as seções aparecem; aqui percorremos os filhos do ``<w:body>``
    para preservar a intercalação real (o que a régua precisa garantir).
    """
    for filho in doc.element.body.iterchildren():
        if filho.tag == qn("w:p"):
            yield Paragraph(filho, doc)
        elif filho.tag == qn("w:tbl"):
            yield Table(filho, doc)


def extrair_docx(caminho: str) -> dict[str, Any]:
    """.docx → ``[estilo, texto]`` de todos os parágrafos e células, em ordem."""
    doc = Document(str(caminho))
    itens: list[list[str]] = []
    for bloco in _iter_blocos(doc):
        if isinstance(bloco, Paragraph):
            itens.append([bloco.style.name, _mascarar(bloco.text)])
        else:  # Table
            for linha in bloco.rows:
                for celula in linha.cells:
                    for par in celula.paragraphs:
                        itens.append([par.style.name, _mascarar(par.text)])
    return {"formato": "docx", "itens": itens}


def extrair_xlsx(caminho: str) -> dict[str, Any]:
    """.xlsx → por aba, ``[coordenada, valor_ou_fórmula]`` das células não vazias."""
    wb = load_workbook(caminho, data_only=False)
    abas: list[dict[str, Any]] = []
    for ws in wb.worksheets:
        celulas: list[list[Any]] = []
        for row in ws.iter_rows():
            for cell in row:
                if cell.value is None:
                    continue
                celulas.append([cell.coordinate, cell.value])
        abas.append({"aba": ws.title, "celulas": celulas})
    return {"formato": "xlsx", "abas": abas}


# ------------------------------------------------------------------ cenários
# Seção de IA idêntica à usada em test_outputs (T-301): exercita o ramo
# "completo" do relatório (rótulo, sumário, roteiro, alertas, confiança).
_SECAO_IA = SecaoIA(
    modo="completo",
    sumario="A dívida do Cartão Banco A é a mais cara e deve ser atacada primeiro.",
    diagnostico="O comprometimento de renda está em nível de atenção.",
    prioridades=["1. Cartão Banco A — maior taxa da carteira"],
    roteiro=[PassoRoteiroIA(credor="Cartão Banco A", abordagem="Quitação à vista",
                            argumentos=["propor desconto para quitação à vista"],
                            concessoes=["usar parte do FGTS"])],
    alertas=["Evite contratar crédito novo enquanto renegocia."],
    confianca=0.85,
    aviso_legal="Este conteúdo é apoio à decisão, não aconselhamento financeiro.",
)

_RUBRICAS = [
    Rubrica("fixas", "contas_casa", "Conta de luz", 180.0, id=1),
    Rubrica("fixas", "contas_casa", "Internet", 120.0, id=2),
    Rubrica("renda", "renda_extra", "Freela", 800.0, id=3),
]

_EVOLUCAO = serie_evolucao([
    ("2026-05", {"variaveis": {"mercado": 750.0}, "fixas": {"moradia": 1400.0}}),
    ("2026-06", {"variaveis": {"mercado": 900.0}, "fixas": {"moradia": 1400.0}}),
])


def _sem_dividas(perfil):
    """Cópia do perfil sem dívidas (não muta a fixture compartilhada)."""
    copia = deepcopy(perfil)
    copia.dividas.clear()
    return copia


# Cada cenário: (formato, construtor(request, caminho) -> caminho_salvo).
# Escolhidos onde o output REALMENTE difere: com/sem dívidas, com/sem
# portabilidade (numeração de seções muda), com/sem seção de IA, e as abas
# opcionais da planilha (rubricas, evolução).
_CENARIOS: dict[str, tuple[str, Callable[[Any, str], str]]] = {
    # .docx — relatório
    "relatorio_atencao": (
        "docx",
        lambda req, c: gerar_relatorio(
            req.getfixturevalue("perfil_atencao"), c,
            extra_mensal=500, nome_usuario="Fulano de Tal"),
    ),
    "relatorio_critico_deficit": (
        "docx",
        lambda req, c: gerar_relatorio(req.getfixturevalue("perfil_critico"), c),
    ),
    "relatorio_saudavel_sem_portabilidade": (
        "docx",
        lambda req, c: gerar_relatorio(
            req.getfixturevalue("perfil_saudavel"), c, extra_mensal=300),
    ),
    "relatorio_atencao_com_ia": (
        "docx",
        lambda req, c: gerar_relatorio(
            req.getfixturevalue("perfil_atencao"), c,
            extra_mensal=500, secao_ia=_SECAO_IA),
    ),
    "relatorio_saudavel_sem_dividas": (
        "docx",
        lambda req, c: gerar_relatorio(
            _sem_dividas(req.getfixturevalue("perfil_saudavel")), c),
    ),
    # .xlsx — planilha
    "planilha_atencao": (
        "xlsx",
        lambda req, c: gerar_planilha(
            req.getfixturevalue("perfil_atencao"), c,
            extra_mensal=500, taxa_alvo_mensal=0.018),
    ),
    "planilha_atencao_rubricas": (
        "xlsx",
        lambda req, c: gerar_planilha(
            req.getfixturevalue("perfil_atencao"), c, rubricas=_RUBRICAS),
    ),
    "planilha_atencao_evolucao": (
        "xlsx",
        lambda req, c: gerar_planilha(
            req.getfixturevalue("perfil_atencao"), c, evolucao=_EVOLUCAO),
    ),
    "planilha_saudavel_sem_dividas": (
        "xlsx",
        lambda req, c: gerar_planilha(
            _sem_dividas(req.getfixturevalue("perfil_saudavel")), c),
    ),
}


# ---------------------------------------------------------------- comparação
def _dumps(dados: Any) -> str:
    """JSON legível em diff: indentado, UTF-8 real, chaves ordenadas.

    ``sort_keys`` ordena apenas as chaves dos dicionários (invólucros
    ``formato``/``itens``/``aba``/``celulas``); as LISTAS — que carregam a ordem
    semântica das seções, células e abas — são preservadas intactas.
    """
    return json.dumps(dados, indent=2, ensure_ascii=False, sort_keys=True) + "\n"


def _comparar_ou_regenerar(nome: str, atual: dict[str, Any]) -> None:
    caminho = GOLDEN_DIR / f"{nome}.json"
    regenerar = os.environ.get("HF_REGENERAR_GOLDEN") == "1"
    em_ci = bool(os.environ.get("CI"))

    if regenerar:
        if em_ci:
            pytest.fail(
                "HF_REGENERAR_GOLDEN=1 com CI setado: regeneração do golden "
                "RECUSADA no pipeline (a régua não se reescreve sozinha). "
                "Regenere localmente, fora de CI, e versione o JSON.")
        GOLDEN_DIR.mkdir(exist_ok=True)
        caminho.write_text(_dumps(atual), encoding="utf-8")
        pytest.skip(f"golden regenerado: {caminho.name}")

    if not caminho.exists():
        pytest.fail(
            f"golden ausente: {caminho}. Gere com HF_REGENERAR_GOLDEN=1 "
            "(fora de CI) e versione o arquivo.")

    esperado = json.loads(caminho.read_text(encoding="utf-8"))
    if atual != esperado:
        diff = "\n".join(difflib.unified_diff(
            _dumps(esperado).splitlines(), _dumps(atual).splitlines(),
            fromfile=f"golden/{nome}.json", tofile=f"{nome} (atual)", lineterm=""))
        pytest.fail(
            f"output divergiu do golden '{nome}' (a refatoração alterou o "
            f"artefato — corrija o código, não o golden):\n{diff}")


@pytest.mark.parametrize("nome", sorted(_CENARIOS))
def test_golden_output(nome: str, request: Any, tmp_path: Path) -> None:
    formato, construir = _CENARIOS[nome]
    caminho = construir(request, str(tmp_path / f"{nome}.{formato}"))
    atual = extrair_docx(caminho) if formato == "docx" else extrair_xlsx(caminho)
    _comparar_ou_regenerar(nome, atual)
