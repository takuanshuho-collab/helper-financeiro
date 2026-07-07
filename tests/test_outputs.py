"""Gate B — geradores de arquivos (H4/REQ-NF-003, H3/REQ-GRD-003; auditoria F-04).

Sem Excel no CI, o "zero erro de fórmula" é verificado estruturalmente:
toda fórmula precisa referenciar abas e células que EXISTEM na pasta gerada
e não pode conter erros literais (#REF!, #NAME?). Complementa: os intervalos
de SOMA precisam cobrir exatamente as linhas de dados.
"""
from __future__ import annotations

import re

from docx import Document
from openpyxl import load_workbook

from contracts import PassoRoteiroIA, SecaoIA
from core.rubricas import Rubrica
from outputs.planilha import gerar_planilha
from outputs.proposta import gerar_proposta
from outputs.relatorio import gerar_relatorio

# Referência de célula, com aba opcional: 'Dívidas'!F4 | D7 | $B$3
_RE_REF = re.compile(r"(?:'([^']+)'!)?\$?([A-Z]{1,3})\$?(\d+)")
_ERROS_LITERAIS = ("#REF!", "#NAME?", "#DIV/0!", "#VALUE!", "#N/A")


def _validar_formulas(caminho) -> list[str]:
    """Retorna a lista de problemas encontrados nas fórmulas. Vazia = Gate B ok."""
    wb = load_workbook(caminho)
    problemas: list[str] = []
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                v = cell.value
                if not (isinstance(v, str) and v.startswith("=")):
                    continue
                for erro in _ERROS_LITERAIS:
                    if erro in v:
                        problemas.append(f"{ws.title}!{cell.coordinate}: contém {erro}")
                for aba, _col, linha in _RE_REF.findall(v):
                    if aba and aba not in wb.sheetnames:
                        problemas.append(
                            f"{ws.title}!{cell.coordinate}: aba inexistente '{aba}' em {v}")
                    if int(linha) < 1:
                        problemas.append(f"{ws.title}!{cell.coordinate}: linha inválida em {v}")
    return problemas


def _texto_docx(caminho) -> str:
    doc = Document(str(caminho))
    partes = [p.text for p in doc.paragraphs]
    for tabela in doc.tables:
        for linha in tabela.rows:
            partes += [c.text for c in linha.cells]
    return "\n".join(partes)


# ------------------------------------------------------------------- planilha
def test_planilha_gate_b_sem_erro_de_formula(perfil_atencao, tmp_path):
    caminho = gerar_planilha(perfil_atencao, str(tmp_path / "diag.xlsx"),
                             extra_mensal=500, taxa_alvo_mensal=0.018)
    assert _validar_formulas(caminho) == []


def test_planilha_com_rubricas_ganha_aba_orcamento(perfil_atencao, tmp_path):
    rubricas = [
        Rubrica("fixas", "contas_casa", "Conta de luz", 180.0, id=1),
        Rubrica("fixas", "contas_casa", "Internet", 120.0, id=2),
        Rubrica("renda", "renda_extra", "Freela", 800.0, id=3),
    ]
    caminho = gerar_planilha(perfil_atencao, str(tmp_path / "diag.xlsx"),
                             rubricas=rubricas)
    wb = load_workbook(caminho)
    assert "Orçamento detalhado" in wb.sheetnames
    ws = wb["Orçamento detalhado"]

    textos = [str(c.value) for row in ws.iter_rows() for c in row if c.value]
    assert any("Conta de luz" in t for t in textos)
    assert any("Freela" in t for t in textos)
    # Subtotal por campo é FÓRMULA (planilha viva), com rótulo pt-BR.
    assert any(t.startswith("Subtotal — Contas da casa") for t in textos)
    assert any(t.startswith("=SUM(") for t in textos), \
        "subtotais das rubricas devem ser fórmulas =SUM"
    # Gate B também vale para a aba nova.
    assert _validar_formulas(caminho) == []


def test_planilha_sem_rubricas_nao_cria_a_aba(perfil_atencao, tmp_path):
    caminho = gerar_planilha(perfil_atencao, str(tmp_path / "diag.xlsx"))
    assert "Orçamento detalhado" not in load_workbook(caminho).sheetnames


def test_planilha_estrutura_e_formulas_chave(perfil_atencao, tmp_path):
    caminho = gerar_planilha(perfil_atencao, str(tmp_path / "diag.xlsx"),
                             extra_mensal=500)
    wb = load_workbook(caminho)
    assert wb.sheetnames[0] == "Diagnóstico"
    assert {"Diagnóstico", "Dívidas", "Estratégias"} <= set(wb.sheetnames)

    ws = wb["Dívidas"]
    n = len(perfil_atencao.dividas)
    primeira, ultima = 4, 4 + n - 1
    # Fórmulas derivadas da 1ª linha de dados: taxa anual e custo total.
    assert ws["E4"].value == "=(1+D4)^12-1"
    assert ws["H4"].value == "=F4*G4"
    # A linha de total soma EXATAMENTE as linhas de dados (nem a mais, nem a menos).
    assert ws.cell(ultima + 1, 3).value == f"=SUM(C{primeira}:C{ultima})"
    # Entradas preservam os valores do perfil.
    assert ws["C4"].value == round(perfil_atencao.dividas[0].saldo_devedor, 2)


def test_planilha_sem_dividas_nao_quebra(perfil_saudavel, tmp_path):
    perfil_saudavel.dividas.clear()
    caminho = gerar_planilha(perfil_saudavel, str(tmp_path / "vazio.xlsx"))
    assert _validar_formulas(caminho) == []


# ------------------------------------------------------------------ relatório
def test_relatorio_conteudo_e_aviso_legal(perfil_atencao, tmp_path):
    caminho = gerar_relatorio(perfil_atencao, str(tmp_path / "rel.docx"),
                              extra_mensal=500, nome_usuario="Fulano de Tal")
    texto = _texto_docx(caminho)
    assert "Relatório de Saúde Financeira" in texto
    assert "Fulano de Tal" in texto
    for d in perfil_atencao.dividas:               # todas as dívidas aparecem
        assert d.credor in texto
    assert "Avalanche" in texto and "Bola de neve" in texto
    assert "apoio à decisão" in texto              # H3: aviso legal presente


def test_relatorio_deficit_gera_alerta(perfil_critico, tmp_path):
    caminho = gerar_relatorio(perfil_critico, str(tmp_path / "rel.docx"))
    texto = _texto_docx(caminho)
    assert "fluxo de caixa está negativo" in texto


# ------------------------------------------- seção "Análise do Agente (IA)"
SECAO_IA = SecaoIA(
    modo="completo",
    sumario="A dívida do Cartão Banco A é a mais cara e deve ser atacada primeiro.",
    diagnostico="O comprometimento de renda está em nível de atenção.",
    prioridades=["1. Cartão Banco A — maior taxa da carteira"],
    roteiro=[PassoRoteiroIA(credor="Cartão Banco A", abordagem="Quitação à vista",
                            argumentos=["propor desconto para quitação à vista"],
                            concessoes=["usar parte do FGTS"])],
    alertas=["Evite contratar crédito novo enquanto renegocia."],
    confianca=0.85,
    aviso_legal="Este conteúdo é apoio à decisão, não aconselhamento financeiro.",
)


def test_relatorio_secao_ia_completa(perfil_atencao, tmp_path):
    """T-301: a narrativa entra em seção própria, separada dos números."""
    caminho = gerar_relatorio(perfil_atencao, str(tmp_path / "rel_ia.docx"),
                              extra_mensal=500, secao_ia=SECAO_IA)
    texto = _texto_docx(caminho)
    assert "Análise do Agente (IA)" in texto
    assert "assistido por IA" in texto             # T-302/P2: rótulo explícito
    assert "fonte oficial" in texto                # REQ-GRD-003: números oficiais = seções determinísticas
    assert "atacada primeiro" in texto             # narrativa presente
    assert "Quitação à vista" in texto             # roteiro presente
    assert "85%" in texto                          # confiança auto-avaliada
    # As seções determinísticas continuam intactas antes da seção de IA.
    assert "Diagnóstico" in texto and "apoio à decisão" in texto


def test_relatorio_secao_ia_degradada_omitida(perfil_atencao, tmp_path):
    """P8/REQ-GRD-003: sem análise aprovada, o .docx sai só com o determinístico."""
    secao = SecaoIA(modo="degradado", motivos=["ERRO_PROVIDER:URLError"])
    caminho = gerar_relatorio(perfil_atencao, str(tmp_path / "rel_deg.docx"),
                              secao_ia=secao)
    texto = _texto_docx(caminho)
    assert "Análise do Agente" not in texto
    assert "Relatório de Saúde Financeira" in texto


# ------------------------------------------------------------------- proposta
def test_proposta_tres_tipos(perfil_atencao, tmp_path):
    divida = perfil_atencao.dividas[0]
    esperado_por_tipo = {
        "quitacao": "QUITAR À VISTA",
        "portabilidade": "PORTABILIDADE",
        "reducao": "RENEGOCIAÇÃO",
    }
    for tipo, esperado in esperado_por_tipo.items():
        caminho = gerar_proposta(divida, str(tmp_path / f"carta_{tipo}.docx"),
                                 tipo=tipo, nome_usuario="Fulano de Tal")
        texto = _texto_docx(caminho)
        assert esperado in texto
        assert divida.credor in texto
        assert "R$ 8.000,00" in texto              # saldo formatado em pt-BR
        assert "Fulano de Tal" in texto


def test_proposta_tipo_desconhecido_cai_em_quitacao(perfil_atencao, tmp_path):
    divida = perfil_atencao.dividas[0]
    caminho = gerar_proposta(divida, str(tmp_path / "carta.docx"), tipo="inexistente")
    assert "QUITAR À VISTA" in _texto_docx(caminho)
