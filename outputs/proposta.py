"""
Geração da carta de proposta de negociação ao credor (.docx).

Três modelos, conforme o objetivo:
  - "quitacao"     : quitação à vista com desconto;
  - "portabilidade": comunicar intenção de portar e pedir contraproposta;
  - "reducao"      : pedir redução de taxa / renegociação das parcelas.
"""
from __future__ import annotations

from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.models import Divida
from core.utils import formatar_brl, formatar_pct

AZUL = RGBColor(0x1F, 0x4E, 0x79)


def _corpo_quitacao(divida, dados):
    valor = dados.get("valor_proposto")
    txt = (
        f"Venho por meio desta manifestar meu interesse em QUITAR À VISTA o "
        f"contrato em referência, cujo saldo devedor informado é de "
        f"{formatar_brl(divida.saldo_devedor)}.\n\n"
        "Considerando a quitação imediata, solicito a apresentação do valor "
        "atualizado com desconto para pagamento único."
    )
    if valor:
        txt += (
            f" Como base para a negociação, proponho o pagamento de "
            f"{formatar_brl(valor)}, à vista, mediante retirada da negativação "
            "e quitação total do contrato."
        )
    return txt


def _corpo_portabilidade(divida, dados):
    taxa_conc = dados.get("taxa_concorrente_mensal")
    banco = dados.get("banco_concorrente", "outra instituição")
    txt = (
        f"Comunico que recebi proposta de PORTABILIDADE de crédito para o contrato "
        f"em referência (saldo devedor de {formatar_brl(divida.saldo_devedor)}), "
        f"junto a {banco}"
    )
    if taxa_conc is not None:
        txt += f", à taxa de {formatar_pct(taxa_conc)} ao mês"
    txt += (
        ".\n\nAntes de efetivar a transferência, solicito que essa instituição "
        "informe se cobre a referida condição ou apresenta proposta mais "
        "vantajosa de redução de taxa, nos termos da regulamentação de "
        "portabilidade de crédito."
    )
    return txt


def _corpo_reducao(divida, dados):
    return (
        f"Venho solicitar a RENEGOCIAÇÃO do contrato em referência, cujo saldo "
        f"devedor é de {formatar_brl(divida.saldo_devedor)} e a taxa atual é de "
        f"{formatar_pct(divida.taxa_mensal)} ao mês.\n\n"
        "Em razão do meu esforço para manter os pagamentos em dia e reorganizar "
        "meu orçamento, solicito a redução da taxa de juros e/ou a readequação do "
        "valor das parcelas a um patamar compatível com minha capacidade de "
        "pagamento, preservando a continuidade do contrato."
    )


_MODELOS = {
    "quitacao": ("Proposta de quitação à vista", _corpo_quitacao),
    "portabilidade": ("Comunicação de portabilidade e pedido de contraproposta", _corpo_portabilidade),
    "reducao": ("Solicitação de renegociação de taxa e parcelas", _corpo_reducao),
}


def gerar_proposta(divida: Divida, caminho_saida: str, tipo: str = "quitacao",
                   dados: dict | None = None, nome_usuario: str = "",
                   cpf: str = "", contrato: str = "") -> str:
    """Monta a carta de proposta e salva em .docx."""
    dados = dados or {}
    if tipo not in _MODELOS:
        tipo = "quitacao"
    titulo, montar_corpo = _MODELOS[tipo]

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Data
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(date.today().strftime("%d de %B de %Y")).font.size = Pt(10)

    # Destinatário
    dest = doc.add_paragraph()
    dest.add_run(f"À {divida.credor}\n").bold = True
    dest.add_run("Setor de Renegociação / Atendimento ao Cliente")

    # Título
    t = doc.add_heading(titulo, level=1)
    for run in t.runs:
        run.font.color.rgb = AZUL

    # Referência do contrato
    ref = doc.add_paragraph()
    partes = []
    if contrato:
        partes.append(f"Contrato nº {contrato}")
    partes.append(f"Modalidade: {divida.tipo}")
    ref.add_run("Ref.: " + " | ".join(partes)).bold = True

    doc.add_paragraph("Prezados,")

    # Corpo (pode ter parágrafos separados por \n\n)
    for bloco in montar_corpo(divida, dados).split("\n\n"):
        doc.add_paragraph(bloco.replace("\n", " "))

    doc.add_paragraph(
        "Solicito que a resposta seja formalizada por escrito, com o valor "
        "atualizado e as condições propostas, para minha análise."
    )
    doc.add_paragraph("Atenciosamente,")

    # Assinatura
    doc.add_paragraph()
    ass = doc.add_paragraph()
    ass.add_run((nome_usuario or "________________________________") + "\n").bold = True
    if cpf:
        ass.add_run(f"CPF: {cpf}")

    doc.save(caminho_saida)
    return caminho_saida
