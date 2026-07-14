"""
Geração do relatório de análise em .docx (python-docx).
"""
from __future__ import annotations

from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from contracts import SecaoIA
from core.diagnostico import resumo_diagnostico
from core.estrategias import comparar_estrategias, gerar_recomendacoes, oportunidades_portabilidade
from core.models import PerfilFinanceiro
from core.utils import formatar_brl, formatar_pct

AZUL = RGBColor(0x1F, 0x4E, 0x79)
CINZA = RGBColor(0x59, 0x59, 0x59)


def _titulo(doc, texto, nivel=1):
    h = doc.add_heading(texto, level=nivel)
    for run in h.runs:
        run.font.color.rgb = AZUL
    return h


def _tabela_indicadores(doc, dados: list[tuple[str, str]]):
    tab = doc.add_table(rows=0, cols=2)
    tab.style = "Light Grid Accent 1"
    for rotulo, valor in dados:
        linha = tab.add_row().cells
        linha[0].text = rotulo
        linha[1].text = valor
        linha[0].paragraphs[0].runs[0].bold = True
    return tab


def _secao_ia(doc, secao: SecaoIA, numero: int) -> None:
    """Renderiza a seção "Análise do Agente (IA)" (T-301, REQ-GRD-003).

    O rótulo de abertura deixa explícito que o conteúdo é assistido por IA e
    que os números oficiais são os das seções determinísticas anteriores (P2).
    """
    _titulo(doc, f"{numero}. Análise do Agente (IA)")

    rotulo = doc.add_paragraph()
    r = rotulo.add_run(
        "Conteúdo assistido por IA (CONSELHEIRO, modelo local): interpretação "
        "gerada exclusivamente a partir dos números das seções anteriores, que "
        "permanecem a fonte oficial deste relatório. Revise antes de agir. "
        f"Confiança auto-avaliada do modelo: {secao.confianca:.0%}."
    )
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = CINZA

    _titulo(doc, "Sumário executivo", nivel=2)
    doc.add_paragraph(secao.sumario)

    _titulo(doc, "Diagnóstico interpretado", nivel=2)
    doc.add_paragraph(secao.diagnostico)

    if secao.prioridades:
        _titulo(doc, "Prioridades sugeridas", nivel=2)
        for prioridade in secao.prioridades:   # já vêm numeradas ("1. Credor — ...")
            doc.add_paragraph(prioridade)

    if secao.roteiro:
        _titulo(doc, "Roteiro de negociação", nivel=2)
        for passo in secao.roteiro:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{passo.credor} — {passo.abordagem}. ").bold = True
            if passo.argumentos:
                p.add_run("Argumentos: " + "; ".join(passo.argumentos) + ". ")
            if passo.concessoes:
                p.add_run("Concessões possíveis: " + "; ".join(passo.concessoes) + ".")

    if secao.alertas:
        _titulo(doc, "Alertas de risco", nivel=2)
        for alerta in secao.alertas:
            doc.add_paragraph("⚠ " + alerta, style="List Bullet")

    if secao.aviso_legal:
        aviso = doc.add_paragraph()
        r = aviso.add_run(secao.aviso_legal)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = CINZA


def _capa(doc, nome_usuario: str) -> None:
    """Renderiza a capa/cabeçalho (título + linha de emissão datada).

    A data vem de ``date.today()``: é o único campo volátil do relatório e por
    isso é mascarada na régua golden — não deve receber tratamento especial aqui.
    """
    cab = doc.add_paragraph()
    cab.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cab.add_run("Relatório de Saúde Financeira")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = AZUL

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    linha_sub = f"Emitido em {date.today().strftime('%d/%m/%Y')}"
    if nome_usuario:
        linha_sub = f"{nome_usuario} — " + linha_sub
    r = sub.add_run(linha_sub)
    r.font.color.rgb = CINZA
    r.font.size = Pt(10)


def _secao_resumo(doc, diag: dict) -> None:
    """Renderiza "1. Resumo executivo" com o alerta de déficit quando aplicável.

    O alerta em vermelho só entra se o fluxo de caixa está negativo, para
    destacar que ele precede qualquer estratégia (P8/leitura do usuário).
    """
    _titulo(doc, "1. Resumo executivo")
    classe = diag["classificacao"]
    p = doc.add_paragraph()
    p.add_run(f"Situação: {classe}. ").bold = True
    p.add_run(diag["classificacao_explicacao"] + " ")
    p.add_run(
        f"O comprometimento de renda com parcelas é de "
        f"{formatar_pct(diag['comprometimento_renda'])}, e o fluxo de caixa mensal "
        f"é de {formatar_brl(diag['fluxo_caixa'])}."
    )
    if diag["tem_deficit"]:
        alerta = doc.add_paragraph()
        run = alerta.add_run(
            "⚠ Atenção: o fluxo de caixa está negativo. As saídas superam as "
            "entradas, o que precisa ser resolvido antes de qualquer estratégia."
        )
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        run.bold = True


def _secao_diagnostico(doc, diag: dict) -> None:
    """Renderiza "2. Diagnóstico" — tabela de indicadores determinísticos (H1)."""
    _titulo(doc, "2. Diagnóstico")
    _tabela_indicadores(doc, [
        ("Renda líquida mensal", formatar_brl(diag["renda_liquida"])),
        ("Despesas totais", formatar_brl(diag["despesas_totais"])),
        ("Total de parcelas/mês", formatar_brl(diag["total_parcelas"])),
        ("Fluxo de caixa (sobra/mês)", formatar_brl(diag["fluxo_caixa"])),
        ("Saldo devedor total", formatar_brl(diag["saldo_devedor_total"])),
        ("Juros futuros embutidos", formatar_brl(diag["juros_totais_futuros"])),
        ("Comprometimento de renda", formatar_pct(diag["comprometimento_renda"])),
    ])


def _secao_ranking(doc, perfil: PerfilFinanceiro, diag: dict) -> None:
    """Renderiza "3. Suas dívidas..." — tabela ordenada do ranking do diagnóstico.

    Sem dívidas cadastradas, degrada para uma linha textual em vez de tabela
    vazia (mantém o relatório coerente para perfis saudáveis, P8).
    """
    _titulo(doc, "3. Suas dívidas, da mais cara para a mais barata")
    if perfil.dividas:
        tab = doc.add_table(rows=1, cols=5)
        tab.style = "Light Grid Accent 1"
        hdr = tab.rows[0].cells
        for i, txt in enumerate(["Credor", "Tipo", "Saldo", "Taxa a.m.", "Parcela"]):
            hdr[i].text = txt
            hdr[i].paragraphs[0].runs[0].bold = True
        for d in diag["ranking"]:
            c = tab.add_row().cells
            c[0].text = d.credor
            c[1].text = d.tipo
            c[2].text = formatar_brl(d.saldo_devedor)
            c[3].text = formatar_pct(d.taxa_mensal)
            c[4].text = formatar_brl(d.parcela)
    else:
        doc.add_paragraph("Nenhuma dívida cadastrada.")


def _secao_estrategia(doc, perfil: PerfilFinanceiro, extra_mensal: float) -> None:
    """Renderiza "4. Estratégia de quitação" (avalanche × bola de neve).

    A nota de transparência do modelo é obrigatória (SPEC REQ-F-003, F-10): a
    simulação é simplificada e serve só para comparar as estratégias entre si.
    """
    _titulo(doc, "4. Estratégia de quitação")
    doc.add_paragraph(
        f"Considerando um pagamento extra de {formatar_brl(extra_mensal)} por mês "
        "além das parcelas mínimas, comparamos os dois métodos clássicos:"
    )
    comp = comparar_estrategias(perfil, extra_mensal)

    def _descreve(nome, res, explicacao):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{nome}: ").bold = True
        if res["quitavel"]:
            p.add_run(
                f"quita em {res['meses']} meses, pagando "
                f"{formatar_brl(res['juros_pagos'])} em juros. {explicacao}"
            )
        else:
            p.add_run(
                "com esse valor extra, as parcelas mínimas não conseguem quitar a "
                "dívida (os juros crescem mais rápido). É preciso aumentar o aporte "
                "ou renegociar as taxas."
            )

    _descreve("Avalanche (ataca o maior juro)", comp["avalanche"],
              "Matematicamente é o caminho mais barato.")
    _descreve("Bola de neve (ataca o menor saldo)", comp["bola_de_neve"],
              "Custa um pouco mais, mas entrega vitórias rápidas que ajudam a manter a disciplina.")

    # Transparência do modelo (SPEC REQ-F-003, auditoria F-10).
    nota = doc.add_paragraph()
    r_nota = nota.add_run(
        "Nota: a simulação usa um modelo simplificado (juros compostos sobre o "
        "saldo, parcela constante) e serve para comparar as estratégias entre "
        "si; o prazo exato pode diferir do cronograma contratual do banco."
    )
    r_nota.italic = True
    r_nota.font.size = Pt(9)
    r_nota.font.color.rgb = CINZA

    if comp["avalanche"]["quitavel"] and comp["bola_de_neve"]["quitavel"]:
        diferenca = comp["bola_de_neve"]["juros_pagos"] - comp["avalanche"]["juros_pagos"]
        if diferenca > 0:
            doc.add_paragraph(
                f"Recomendação: a avalanche economiza {formatar_brl(diferenca)} em "
                "juros. Prefira-a, a menos que você precise do impulso psicológico "
                "das quitações rápidas da bola de neve."
            )


def _secao_portabilidade(doc, ops: list, taxa_alvo_mensal: float) -> None:
    """Renderiza "5. Oportunidades de portabilidade".

    Só é chamada quando há oportunidades (``ops`` não vazio); a ausência dela
    desloca a numeração das seções seguintes — comportamento fixado no golden
    ``relatorio_saudavel_sem_portabilidade``.
    """
    _titulo(doc, "5. Oportunidades de portabilidade")
    doc.add_paragraph(
        f"Migrando as dívidas caras para uma taxa de "
        f"{formatar_pct(taxa_alvo_mensal)} a.m., a economia estimada seria:"
    )
    tab = doc.add_table(rows=1, cols=3)
    tab.style = "Light Grid Accent 1"
    hdr = tab.rows[0].cells
    for i, txt in enumerate(["Credor", "Economia mensal", "Economia total"]):
        hdr[i].text = txt
        hdr[i].paragraphs[0].runs[0].bold = True
    for o in ops:
        c = tab.add_row().cells
        c[0].text = o["credor"]
        c[1].text = formatar_brl(o["economia_mensal"])
        c[2].text = formatar_brl(o["economia_total"])


def _secao_recomendacoes(doc, perfil: PerfilFinanceiro, diag: dict, numero: int) -> None:
    """Renderiza a seção "Recomendações" com número dinâmico (5 ou 6).

    O ``numero`` é calculado pelo chamador porque depende da presença da seção
    de portabilidade, que pode não existir.
    """
    _titulo(doc, f"{numero}. Recomendações")
    for rec in gerar_recomendacoes(perfil, diag):
        doc.add_paragraph(rec, style="List Bullet")


def _secao_proximos_passos(doc, numero: int) -> None:
    """Renderiza a seção "Próximos passos" com número dinâmico (6 ou 7).

    Lista fixa de ações; o ``numero`` acompanha o deslocamento da portabilidade.
    """
    _titulo(doc, f"{numero}. Próximos passos")
    for passo in [
        "Solicite a cada credor o saldo devedor atualizado, por escrito.",
        "Simule a portabilidade nos concorrentes e use as propostas como alavanca.",
        "Negocie primeiro as dívidas mais caras (maior taxa/CET).",
        "Registre todo acordo por escrito (protocolo ou Consumidor.gov.br).",
    ]:
        doc.add_paragraph(passo, style="List Number")


def _aviso_final(doc) -> None:
    """Renderiza o parágrafo de aviso legal que encerra o relatório."""
    doc.add_paragraph()
    aviso = doc.add_paragraph()
    r = aviso.add_run(
        "Aviso: este relatório é uma ferramenta de apoio à decisão baseada nos "
        "dados informados. Não constitui aconselhamento financeiro ou de "
        "investimento personalizado. As taxas de mercado e as regras de programas "
        "de renegociação mudam; confirme os números vigentes antes de decidir."
    )
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = CINZA


def gerar_relatorio(perfil: PerfilFinanceiro, caminho_saida: str,
                    extra_mensal: float = 0.0, taxa_alvo_mensal: float = 0.018,
                    nome_usuario: str = "",
                    secao_ia: SecaoIA | None = None) -> str:
    diag = resumo_diagnostico(perfil)
    doc = Document()

    # Fonte padrão
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _capa(doc, nome_usuario)
    _secao_resumo(doc, diag)
    _secao_diagnostico(doc, diag)
    _secao_ranking(doc, perfil, diag)
    _secao_estrategia(doc, perfil, extra_mensal)

    # A portabilidade só entra quando há oportunidade; sua presença desloca a
    # numeração das seções seguintes (fixada no golden sem portabilidade).
    ops = oportunidades_portabilidade(perfil, taxa_alvo_mensal)
    if ops:
        _secao_portabilidade(doc, ops, taxa_alvo_mensal)

    _secao_recomendacoes(doc, perfil, diag, numero=6 if ops else 5)
    _secao_proximos_passos(doc, numero=7 if ops else 6)

    # --- Análise do Agente (IA) — T-301 ---
    # Seção SEPARADA das tabelas determinísticas (REQ-GRD-003): só entra quando
    # a análise passou por todos os guardrails (modo "completo"); em modo
    # degradado o relatório sai apenas com o determinístico (P8).
    if secao_ia is not None and secao_ia.modo == "completo":
        _secao_ia(doc, secao_ia, numero=8 if ops else 7)

    _aviso_final(doc)

    doc.save(caminho_saida)
    return caminho_saida
